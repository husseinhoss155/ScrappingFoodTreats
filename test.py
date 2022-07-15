import os
from docx import Document
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
driver = webdriver.Chrome('chromedriver.exe')
driver.get('https://closelyhq.com/templates/')

#Waiting for a random element to load for preventing getting error loading
element = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located(
                (By.XPATH, '//*[@id="__next"]/div/section[2]/div[1]/div[2]/div[1]/div[2]/div'))
        )

#The number of connection messages in 7 so the first 7 categories are the connection messages
counter=1

#Initializing the output folders
os.makedirs('Connection messages')
os.makedirs('Goals')

#Useful for the XPATH
i = 2

for div in driver.find_elements(By.CLASS_NAME,'NavigationItem_wrapper__Xy2GH'):
    #Connection Messages
    if counter<=7:
        #Click on the category
        div.click()

        #Scraping the title and removing the number at the beggining of the title
        title=div.text[3:]

        #For the scraped templates and sequences
        texts=[]
        document = Document()

        #Iterating over the whole texts
        for div in driver.find_elements(By.CLASS_NAME,'Sequence_wrapper__rG3Ru'):
            texts.append(div.text)
            texts.append('\n\n')

        paragraph = document.add_paragraph('')

        #This loop is for making the specified words bold
        for template in texts:
            lines = template.splitlines()
            for line in lines:
                if 'Template' in line or 'Sequence' in line or 'Recommended' in line:
                    to_be_bold = paragraph.add_run(line)
                    to_be_bold.bold = True
                    paragraph.add_run('\n')
                else:
                    try:
                        int(line.split(' ')[0])
                        number_to_bold = paragraph.add_run(line)
                        number_to_bold.bold = True
                        paragraph.add_run('\n')
                    except:
                        paragraph.add_run(line)
                        paragraph.add_run('\n')
        document.save('Connection messages/'+title+'.docx')
        counter+=1
    else:
        #Goals
        xpath='//*[@id="__next"]/div/section[2]/div[1]/div[2]/div[1]/div[4]/div/div['+str(i)+']'
        div.click()

        #Scraping the folder name and removing the number at the beggining of the title
        folder = div.text[3:]
        if '/' in folder:
            folder = folder.replace('/', '_')
        os.makedirs('Goals/' + folder)

        #Useful for XPATH
        i+=2
        j = 1

        #Iterating over each subcategory from each category
        while True:
            try:
                xpathh = xpath + '/div['+str(j)+']/div[2]'
                j+=1
                div = driver.find_element(By.XPATH,xpathh)
                if div.text=='':
                    continue

                #Scraping the title (no number at the beggining)
                title = div.text
                div.click()

                # For the scraped templates and sequences
                texts = []
                document = Document()

                # Iterating over the whole texts
                for div in driver.find_elements(By.CLASS_NAME, 'Sequence_wrapper__rG3Ru'):
                    texts.append(div.text)
                    texts.append('\n\n')

                # This loop is for making the specified words bold
                paragraph = document.add_paragraph('')
                for template in texts:
                    lines = template.splitlines()
                    for line in lines:
                        if 'Template' in line or 'Sequence' in line or 'Recommended' in line:
                            to_be_bold = paragraph.add_run(line)
                            to_be_bold.bold = True
                            paragraph.add_run('\n')
                        else:
                            try:
                                int(line.split(' ')[0])
                                number_to_bold = paragraph.add_run(line)
                                number_to_bold.bold = True
                                paragraph.add_run('\n')
                            except:
                                paragraph.add_run(line)
                                paragraph.add_run('\n')
                document.save('Goals/' +folder+'/'+ title + '.docx')
                counter += 1
            except:
                break

#Opening the project folder to check the outputs
os.startfile('')

#Closing the driver
driver.quit()